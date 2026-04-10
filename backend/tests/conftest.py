"""
JurisIA — Configuration Pytest Globale
Fixtures partagées entre tous les tests.
Ce fichier est automatiquement chargé par pytest.
"""
import logging
import os
import pytest
import structlog

# ── Configurer structlog immédiatement au niveau module ───────────────────────
# DOIT être fait avant tout import de app.* pour éviter le crash PrintLogger.name
structlog.configure(
    processors=[
        structlog.stdlib.add_log_level,
        structlog.dev.ConsoleRenderer(colors=False),
    ],
    wrapper_class=structlog.make_filtering_bound_logger(logging.WARNING),
    logger_factory=structlog.stdlib.LoggerFactory(),
    cache_logger_on_first_use=False,
)

# ── Variables d'env de test (aussi au niveau module) ─────────────────────────
_TEST_ENV = {
    "APP_ENV":               "development",
    "APP_DEBUG":             "false",
    "APP_SECRET_KEY":        "test_secret_key_32_chars_minimum_ok!!",
    "JWT_SECRET_KEY":        "test_jwt_64_chars_min_required_for_hmac_sha256_signing_ok",
    "ENCRYPTION_KEY":        "oudIUTdrF0o6QrztsCaqJCKd8L2f5au88JBW-WzeKgw=",
    "DATABASE_URL":          "postgresql+asyncpg://test:test@localhost:5432/jurisai_test",
    "REDIS_URL":             "redis://localhost:6379/15",
    "MISTRAL_API_KEY":       "dummy_mistral_key_for_tests",
    "STRIPE_SECRET_KEY":     "sk_test_dummy_stripe_key",
    "ALLOWED_ORIGINS":       "http://localhost:3000",
    "CELERY_BROKER_URL":     "redis://localhost:6379/14",
    "CELERY_RESULT_BACKEND": "redis://localhost:6379/14",
}
os.environ.update(_TEST_ENV)

# Clear le cache settings après avoir défini les vars
from app.core.config import get_settings  # noqa: E402
get_settings.cache_clear()


@pytest.fixture(autouse=True)
def reset_structlog():
    """Réinitialise structlog pour chaque test (évite pollution entre tests)."""
    structlog.configure(
        processors=[
            structlog.stdlib.add_log_level,
            structlog.dev.ConsoleRenderer(colors=False),
        ],
        wrapper_class=structlog.make_filtering_bound_logger(logging.WARNING),
        logger_factory=structlog.stdlib.LoggerFactory(),
        cache_logger_on_first_use=False,
    )
    yield


@pytest.fixture
def sample_user_data():
    return {
        "id":             "01TESTUSER000000000000000A",
        "email":          "marie.dupont@entreprise.fr",
        "full_name":      "Marie Dupont",
        "email_verified": True,
        "two_fa_enabled": False,
    }


@pytest.fixture
def sample_org_data():
    return {
        "id":                   "01TESTORG0000000000000000",
        "name":                 "Dupont & Associes SAS",
        "siren":                "123456789",
        "sector_label":         "Conseil et services",
        "employee_count_range": "11-50",
    }


@pytest.fixture
def sample_contract_text():
    return """
CONTRAT DE PRESTATION DE SERVICES

Article 1 : Objet
Le Prestataire s'engage a fournir des services de conseil.

Article 2 : Prix
Les services sont factures au tarif de 1.500 euros HT par jour.

Article 3 : Responsabilite
Le Prestataire sera responsable de tout prejudice direct ou indirect
cause au Client dans le cadre de l'execution du present contrat.

Article 4 : Droit applicable
Le present contrat est soumis au droit francais.
    """


@pytest.fixture
def valid_pdf_bytes():
    pdf = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>>>endobj
xref
0 4
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
trailer<</Size 4/Root 1 0 R>>
startxref
217
%%EOF"""
    return pdf


@pytest.fixture
def valid_docx_bytes():
    import io
    from docx import Document
    doc = Document()
    doc.add_heading("Test Contrat", level=1)
    doc.add_paragraph("Contrat de test pour JurisIA.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
