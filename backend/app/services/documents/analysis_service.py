"""
JurisIA — Service d'Analyse de Documents (Cœur Métier)
Pipeline RAG : parsing → chunking → embedding → retrieval → analyse Mistral → résultat structuré.
C'est la fonctionnalité #1 qui crée la valeur principale du produit.
"""
from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from typing import Optional

import structlog
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.security import generate_ulid
from app.models import Document, DocumentClause, DocumentStatus, DocumentType, RiskLevel

logger = structlog.get_logger(__name__)


# ─── Structures de données ────────────────────────────────────────────────────

@dataclass
class ParsedDocument:
    """Document parsé avec son contenu textuel extrait."""
    text: str
    page_count: int
    file_size_bytes: int
    detected_type: str  # "contract", "cgv", "rh", "other"
    language: str       # "fr", "en", etc.


@dataclass
class ClauseAnalysis:
    """Résultat d'analyse d'une clause individuelle."""
    clause_text: str
    risk_level: RiskLevel
    explanation: str
    suggestion: Optional[str]
    legal_reference: Optional[str]
    legal_reference_url: Optional[str]
    position_start: Optional[int]
    position_end: Optional[int]


@dataclass
class DocumentAnalysisResult:
    """Résultat complet de l'analyse d'un document."""
    score: int                          # 0-100
    clauses: list[ClauseAnalysis]
    summary: str                        # Résumé exécutif < 200 mots
    document_type_detected: str
    total_clauses: int
    risk_counts: dict[str, int]         # {"danger": 2, "warning": 3, "safe": 8}


# ─── Prompt Engineering ───────────────────────────────────────────────────────

ANALYSIS_SYSTEM_PROMPT = """Tu es JurisIA, un assistant juridique expert en droit français des affaires et droit du travail.
Ta mission est d'analyser des documents juridiques et d'identifier les clauses à risque pour des dirigeants de PME non-juristes.

RÈGLES ABSOLUES :
1. Réponds UNIQUEMENT en JSON valide selon le schéma fourni
2. Utilise un langage accessible, jamais de jargon juridique sans explication
3. Cite TOUJOURS la référence légale applicable (Code civil, Code du travail, Légifrance)
4. Sois CONSERVATEUR dans l'évaluation des risques : vaut mieux sur-alerter que sous-alerter
5. Ne génère JAMAIS de conseil juridique direct - tu fournis une analyse d'aide à la décision
6. Si tu n'es pas certain, indique le niveau de confiance dans l'explication

NIVEAUX DE RISQUE :
- "danger" : clause pouvant causer un préjudice financier ou juridique significatif
- "warning" : clause à surveiller, à négocier ou à clarifier
- "safe" : clause conforme et équilibrée
- "missing" : clause importante qui devrait être présente mais est absente

RÉFÉRENCES LÉGALES :
- Toujours citer sous la forme "Article L1234-5 du Code du travail" ou "Article 1234 du Code civil"
- URL format : "https://www.legifrance.gouv.fr/codes/article_lc/LEGIARTI..."
"""

ANALYSIS_USER_PROMPT_TEMPLATE = """Analyse ce document juridique et retourne un JSON structuré.

DOCUMENT À ANALYSER :
```
{document_text}
```

CONTEXTE ENTREPRISE :
- Secteur : {sector}
- Taille : {employee_count}

SCHÉMA JSON ATTENDU (retourne UNIQUEMENT ce JSON, sans markdown) :
{{
  "document_type": "contract|cgv|rh_contract|mise_en_demeure|other",
  "summary": "Résumé en 3-4 phrases du document et de ses enjeux principaux",
  "score": <entier 0-100, 100=parfait>,
  "clauses": [
    {{
      "clause_text": "Texte exact ou paraphrase de la clause",
      "risk_level": "danger|warning|safe|missing",
      "explanation": "Explication en français simple du problème ou de la conformité",
      "suggestion": "Action concrète recommandée (null si safe)",
      "legal_reference": "Article X du Code Y (null si non applicable)",
      "legal_reference_url": "URL Légifrance (null si inconnu)",
      "position_approximate": <numéro de clause dans le document, null si missing>
    }}
  ]
}}

IMPORTANT : 
- Analyse TOUTES les clauses importantes, même les conformes (pour rassurer l'utilisateur)
- Pour les clauses manquantes importantes, crée une entrée avec risk_level="missing"
- Le score doit refléter l'équilibre global : 100 = contrat parfaitement équilibré et complet
"""


# ─── Service Principal ────────────────────────────────────────────────────────

class DocumentAnalysisService:
    """
    Service d'analyse de documents juridiques.
    Orchestre le pipeline : parsing → préparation → appel Mistral → post-traitement → sauvegarde.
    """

    def __init__(self, db: AsyncSession):
        self.db = db

    async def analyze_document(
        self,
        document_id: str,
        file_content: bytes,
        filename: str,
        organization_sector: str = "services",
        employee_count: str = "11-50",
    ) -> DocumentAnalysisResult:
        """
        Analyse complète d'un document. Méthode principale du service.

        Args:
            document_id: ID du document en BDD (déjà créé avec status=PENDING)
            file_content: Contenu binaire du fichier
            filename: Nom du fichier (pour détecter le format)
            organization_sector: Secteur de l'entreprise
            employee_count: Taille de l'entreprise

        Returns:
            DocumentAnalysisResult avec toutes les clauses analysées
        """
        logger.info("document_analysis_started", document_id=document_id, filename=filename)

        # 1. Mettre à jour le statut
        await self._update_status(document_id, DocumentStatus.PROCESSING)

        try:
            # 2. Parser le document
            parsed = await self._parse_document(file_content, filename)
            logger.info(
                "document_parsed",
                document_id=document_id,
                pages=parsed.page_count,
                chars=len(parsed.text),
            )

            # 3. Préparer le texte (tronquer si trop long pour le contexte LLM)
            prepared_text = self._prepare_text_for_llm(parsed.text)

            # 4. Appeler Mistral pour l'analyse
            raw_result = await self._call_mistral_analysis(
                document_text=prepared_text,
                sector=organization_sector,
                employee_count=employee_count,
            )

            # 5. Parser et valider le JSON de réponse
            analysis_result = self._parse_mistral_response(raw_result)

            # 6. Sauvegarder les résultats en BDD
            await self._save_analysis_results(
                document_id=document_id,
                parsed=parsed,
                result=analysis_result,
            )

            logger.info(
                "document_analysis_completed",
                document_id=document_id,
                score=analysis_result.score,
                clauses_count=len(analysis_result.clauses),
            )

            return analysis_result

        except Exception as e:
            logger.error("document_analysis_failed", document_id=document_id, error=str(e))
            await self._update_status(document_id, DocumentStatus.FAILED, error_message=str(e))
            raise

    # ── Parsing ────────────────────────────────────────────────────────────

    async def _parse_document(self, content: bytes, filename: str) -> ParsedDocument:
        """Extrait le texte du document selon son format."""
        filename_lower = filename.lower()
        text = ""
        page_count = 1

        if filename_lower.endswith(".pdf"):
            text, page_count = self._parse_pdf(content)
        elif filename_lower.endswith((".docx", ".doc")):
            text = self._parse_docx(content)
        elif filename_lower.endswith(".txt"):
            text = content.decode("utf-8", errors="replace")
        else:
            raise ValueError(f"Format de fichier non supporté : {filename}. Formats acceptés : PDF, DOCX, TXT")

        if not text.strip():
            raise ValueError("Le document semble vide ou ne contient pas de texte extractible")

        detected_type = self._detect_document_type(text)

        return ParsedDocument(
            text=text,
            page_count=page_count,
            file_size_bytes=len(content),
            detected_type=detected_type,
            language="fr",
        )

    def _parse_pdf(self, content: bytes) -> tuple[str, int]:
        """Parse un PDF et retourne (texte, nombre de pages)."""
        import pdfplumber
        text_parts = []
        page_count = 0

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return "\n\n".join(text_parts), page_count

    def _parse_docx(self, content: bytes) -> str:
        """Parse un fichier DOCX et retourne le texte."""
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _detect_document_type(self, text: str) -> str:
        """Détecte automatiquement le type de document."""
        text_lower = text.lower()

        patterns = {
            "contract": ["contrat de prestation", "contrat de service", "accord entre"],
            "cgv": ["conditions générales de vente", "cgv", "conditions générales"],
            "nda": ["accord de confidentialité", "non-disclosure", "secret professionnel"],
            "rh_contract": ["contrat de travail", "contrat à durée", "contrat d'apprentissage"],
            "mise_en_demeure": ["mise en demeure", "sommation", "injonction"],
            "bail": ["contrat de bail", "bailleur", "preneur"],
        }

        for doc_type, keywords in patterns.items():
            if any(kw in text_lower for kw in keywords):
                return doc_type

        return "other"

    # ── Préparation du texte ───────────────────────────────────────────────

    def _prepare_text_for_llm(self, text: str, max_chars: int = 12000) -> str:
        """
        Prépare le texte pour l'envoi au LLM.
        Limite la taille pour ne pas dépasser la fenêtre de contexte.
        Pour les documents longs, garde le début et la fin (les clauses importantes).
        """
        if len(text) <= max_chars:
            return text

        # Garder 60% du début et 40% de la fin pour les documents longs
        start_chars = int(max_chars * 0.6)
        end_chars = max_chars - start_chars

        start_text = text[:start_chars]
        end_text = text[-end_chars:]

        truncation_notice = f"\n\n[... {len(text) - max_chars} caractères omis pour analyse ...]\n\n"

        logger.warning(
            "document_truncated_for_llm",
            original_chars=len(text),
            truncated_to=max_chars,
        )

        return start_text + truncation_notice + end_text

    # ── Appel Mistral ──────────────────────────────────────────────────────

    async def _call_mistral_analysis(
        self,
        document_text: str,
        sector: str,
        employee_count: str,
    ) -> str:
        """
        Appelle l'API Mistral pour analyser le document.
        Retourne la réponse brute JSON.
        """
        from mistralai import Mistral

        client = Mistral(api_key=settings.MISTRAL_API_KEY)

        user_prompt = ANALYSIS_USER_PROMPT_TEMPLATE.format(
            document_text=document_text,
            sector=sector,
            employee_count=employee_count,
        )

        # Retry logic : 3 tentatives avec backoff exponentiel
        from tenacity import retry, stop_after_attempt, wait_exponential
        import asyncio

        for attempt in range(3):
            try:
                response = client.chat.complete(
                    model=settings.MISTRAL_MODEL,
                    messages=[
                        {"role": "system", "content": ANALYSIS_SYSTEM_PROMPT},
                        {"role": "user", "content": user_prompt},
                    ],
                    max_tokens=settings.MISTRAL_MAX_TOKENS,
                    temperature=settings.MISTRAL_TEMPERATURE,
                )
                return response.choices[0].message.content

            except Exception as e:
                if attempt == 2:
                    raise RuntimeError(f"Échec de l'analyse IA après 3 tentatives : {e}") from e
                wait_time = 2 ** attempt
                logger.warning(f"Mistral API attempt {attempt + 1} failed, retrying in {wait_time}s: {e}")
                await asyncio.sleep(wait_time)

    # ── Post-traitement ────────────────────────────────────────────────────

    def _parse_mistral_response(self, raw_response: str) -> DocumentAnalysisResult:
        """
        Parse et valide la réponse JSON de Mistral.
        Robuste aux variations de format (JSON partiel, commentaires, etc.).
        """
        # Nettoyer les éventuels backticks markdown
        clean = raw_response.strip()
        clean = re.sub(r"^```json\s*", "", clean, flags=re.MULTILINE)
        clean = re.sub(r"```\s*$", "", clean, flags=re.MULTILINE)
        clean = clean.strip()

        try:
            data = json.loads(clean)
        except json.JSONDecodeError as e:
            # Tentative de récupération : extraire le JSON entre les premières { }
            match = re.search(r"\{.*\}", clean, re.DOTALL)
            if match:
                try:
                    data = json.loads(match.group())
                except json.JSONDecodeError:
                    raise ValueError(f"Impossible de parser la réponse IA : {e}")
            else:
                raise ValueError(f"La réponse IA n'est pas du JSON valide : {e}")

        # Construire les clauses typées
        clauses = []
        for raw_clause in data.get("clauses", []):
            risk_str = raw_clause.get("risk_level", "warning").lower()
            try:
                risk = RiskLevel(risk_str)
            except ValueError:
                risk = RiskLevel.WARNING

            clause = ClauseAnalysis(
                clause_text=str(raw_clause.get("clause_text", ""))[:2000],
                risk_level=risk,
                explanation=str(raw_clause.get("explanation", ""))[:1000],
                suggestion=raw_clause.get("suggestion"),
                legal_reference=raw_clause.get("legal_reference"),
                legal_reference_url=raw_clause.get("legal_reference_url"),
                position_start=raw_clause.get("position_approximate"),
                position_end=None,
            )
            clauses.append(clause)

        # Calculer le score et les compteurs de risque
        risk_counts = {"danger": 0, "warning": 0, "safe": 0, "missing": 0}
        for clause in clauses:
            risk_counts[clause.risk_level.value] += 1

        # Valider/corriger le score
        raw_score = data.get("score", 50)
        score = max(0, min(100, int(raw_score)))

        return DocumentAnalysisResult(
            score=score,
            clauses=clauses,
            summary=str(data.get("summary", "Analyse complétée."))[:500],
            document_type_detected=str(data.get("document_type", "other")),
            total_clauses=len(clauses),
            risk_counts=risk_counts,
        )

    # ── Sauvegarde BDD ─────────────────────────────────────────────────────

    async def _save_analysis_results(
        self,
        document_id: str,
        parsed: ParsedDocument,
        result: DocumentAnalysisResult,
    ) -> None:
        """Sauvegarde les résultats d'analyse en base de données."""
        document = await self.db.get(Document, document_id)
        if not document:
            return

        # Mettre à jour le document
        document.status = DocumentStatus.COMPLETED
        document.score = result.score
        document.analysis_result = {
            "summary": result.summary,
            "document_type_detected": result.document_type_detected,
            "total_clauses": result.total_clauses,
            "risk_counts": result.risk_counts,
        }
        document.metadata_ = {
            "page_count": parsed.page_count,
            "file_size_bytes": parsed.file_size_bytes,
            "detected_type": parsed.detected_type,
            "language": parsed.language,
        }

        # Sauvegarder chaque clause
        for clause in result.clauses:
            db_clause = DocumentClause(
                id=generate_ulid(),
                document_id=document_id,
                clause_text=clause.clause_text,
                risk_level=clause.risk_level,
                explanation=clause.explanation,
                suggestion=clause.suggestion,
                legal_reference=clause.legal_reference,
                legal_reference_url=clause.legal_reference_url,
                position_start=clause.position_start,
                position_end=clause.position_end,
            )
            self.db.add(db_clause)

        await self.db.flush()

    async def _update_status(
        self,
        document_id: str,
        new_status: DocumentStatus,
        error_message: Optional[str] = None,
    ) -> None:
        """Met à jour le statut d'un document."""
        document = await self.db.get(Document, document_id)
        if document:
            document.status = new_status
            if error_message:
                document.error_message = error_message[:1000]
            await self.db.flush()
