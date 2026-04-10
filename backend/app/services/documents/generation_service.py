"""
JurisIA — Service de Génération de Documents Juridiques
Produit des documents légaux personnalisés via LLM + templates structurés.
Chaque template est versionné et daté pour la traçabilité légale.
"""
from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Optional

import structlog
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models import Document, DocumentStatus

logger = structlog.get_logger(__name__)

# ─── Disclaimer légal inséré dans chaque document généré ─────────────────────
LEGAL_DISCLAIMER = (
    "⚠️ IMPORTANT — Ce document a été généré par JurisIA (IA) à titre d'aide à la rédaction. "
    "Il ne constitue pas un conseil juridique professionnel. "
    "Pour tout enjeu financier supérieur à 10 000€ ou toute situation complexe (licenciement, litige), "
    "consultez un avocat. Généré le {date} | Version template : {version} | JurisIA v1.0"
)

# ─── Catalogue de templates ───────────────────────────────────────────────────
DOCUMENT_TEMPLATES = {
    "prestation_services": {
        "name": "Contrat de Prestation de Services",
        "category": "contract",
        "description": "Contrat standard pour les missions de conseil, développement, formation, etc.",
        "version": "1.2",
        "required_plan": "free",
        "fields": [
            {"key": "prestataire_nom", "label": "Nom du prestataire", "type": "text", "required": True},
            {"key": "prestataire_siren", "label": "SIREN prestataire", "type": "text", "required": False},
            {"key": "client_nom", "label": "Nom du client", "type": "text", "required": True},
            {"key": "client_siren", "label": "SIREN client", "type": "text", "required": False},
            {"key": "objet_mission", "label": "Objet de la mission", "type": "textarea", "required": True},
            {"key": "date_debut", "label": "Date de début", "type": "date", "required": True},
            {"key": "date_fin", "label": "Date de fin (ou durée)", "type": "text", "required": False},
            {"key": "montant_ht", "label": "Montant HT (€)", "type": "number", "required": True},
            {"key": "modalite_paiement", "label": "Modalité de paiement", "type": "select", "required": True,
             "options": ["30 jours fin de mois", "À réception de facture", "Acompte 30% / solde à livraison", "Mensuel"]},
            {"key": "lieu_execution", "label": "Lieu d'exécution", "type": "text", "required": False},
            {"key": "confidentialite", "label": "Clause de confidentialité ?", "type": "boolean", "required": True},
        ],
        "prompt_template": "prestation_services_v1",
    },
    "cgv": {
        "name": "Conditions Générales de Vente (CGV)",
        "category": "contract",
        "description": "CGV conformes au droit français pour vente de produits ou services B2B/B2C.",
        "version": "1.1",
        "required_plan": "free",
        "fields": [
            {"key": "entreprise_nom", "label": "Nom de l'entreprise", "type": "text", "required": True},
            {"key": "entreprise_siren", "label": "SIREN", "type": "text", "required": True},
            {"key": "type_activite", "label": "Type d'activité", "type": "select", "required": True,
             "options": ["Vente de produits B2B", "Vente de produits B2C", "Prestations de services B2B", "SaaS/Logiciel"]},
            {"key": "prix_ttc", "label": "Les prix sont affichés TTC ?", "type": "boolean", "required": True},
            {"key": "delai_livraison", "label": "Délai de livraison standard", "type": "text", "required": False},
            {"key": "politique_retour", "label": "Politique de retour/remboursement", "type": "textarea", "required": False},
            {"key": "juridiction", "label": "Tribunal compétent en cas de litige", "type": "text", "required": True,
             "placeholder": "Ex: Tribunal de Commerce de Paris"},
        ],
        "prompt_template": "cgv_v1",
    },
    "nda": {
        "name": "Accord de Confidentialité (NDA)",
        "category": "contract",
        "description": "NDA bilatéral ou unilatéral, adapté au contexte commercial français.",
        "version": "1.0",
        "required_plan": "free",
        "fields": [
            {"key": "partie_1_nom", "label": "Nom de la 1ère partie", "type": "text", "required": True},
            {"key": "partie_2_nom", "label": "Nom de la 2ème partie", "type": "text", "required": True},
            {"key": "type_nda", "label": "Type de NDA", "type": "select", "required": True,
             "options": ["Bilatéral (les deux parties partagent des infos)", "Unilatéral (une seule partie divulgue)"]},
            {"key": "objet_partage", "label": "Nature des informations confidentielles", "type": "textarea", "required": True},
            {"key": "duree_annees", "label": "Durée de la confidentialité (années)", "type": "number", "required": True},
            {"key": "contexte", "label": "Contexte (levée de fonds, partenariat, recrutement...)", "type": "text", "required": False},
        ],
        "prompt_template": "nda_v1",
    },
    "lettre_relance_1": {
        "name": "Lettre de Relance Amiable (Niveau 1)",
        "category": "recovery",
        "description": "Première relance courtoise pour facture impayée.",
        "version": "1.0",
        "required_plan": "free",
        "fields": [
            {"key": "expediteur_nom", "label": "Votre nom/société", "type": "text", "required": True},
            {"key": "destinataire_nom", "label": "Nom du débiteur", "type": "text", "required": True},
            {"key": "numero_facture", "label": "Numéro de facture", "type": "text", "required": True},
            {"key": "montant_ttc", "label": "Montant TTC dû (€)", "type": "number", "required": True},
            {"key": "date_echeance", "label": "Date d'échéance initiale", "type": "date", "required": True},
            {"key": "iban", "label": "IBAN pour règlement", "type": "text", "required": False},
        ],
        "prompt_template": "relance_1_v1",
    },
    "mise_en_demeure": {
        "name": "Mise en Demeure",
        "category": "recovery",
        "description": "Mise en demeure formelle avant action judiciaire. Dernière étape amiable.",
        "version": "1.0",
        "required_plan": "starter",
        "fields": [
            {"key": "expediteur_nom", "label": "Votre nom/société", "type": "text", "required": True},
            {"key": "expediteur_adresse", "label": "Votre adresse complète", "type": "textarea", "required": True},
            {"key": "destinataire_nom", "label": "Nom du destinataire", "type": "text", "required": True},
            {"key": "destinataire_adresse", "label": "Adresse du destinataire", "type": "textarea", "required": True},
            {"key": "objet_demande", "label": "Ce que vous réclamez (paiement, exécution...)", "type": "textarea", "required": True},
            {"key": "montant", "label": "Montant réclamé (€)", "type": "number", "required": False},
            {"key": "delai_reponse_jours", "label": "Délai pour répondre (jours)", "type": "number", "required": True},
            {"key": "reference_legale", "label": "Contrat/facture de référence", "type": "text", "required": False},
        ],
        "prompt_template": "mise_en_demeure_v1",
    },
    "cdi": {
        "name": "Contrat de Travail CDI",
        "category": "rh",
        "description": "Contrat CDI conforme au Code du travail français.",
        "version": "1.0",
        "required_plan": "pro",
        "fields": [
            {"key": "employeur_nom", "label": "Nom de l'employeur", "type": "text", "required": True},
            {"key": "employe_nom", "label": "Nom du salarié", "type": "text", "required": True},
            {"key": "poste", "label": "Intitulé du poste", "type": "text", "required": True},
            {"key": "date_embauche", "label": "Date d'embauche", "type": "date", "required": True},
            {"key": "salaire_brut", "label": "Salaire brut mensuel (€)", "type": "number", "required": True},
            {"key": "temps_travail_heures", "label": "Durée hebdomadaire (heures)", "type": "number", "required": True},
            {"key": "periode_essai_mois", "label": "Période d'essai (mois, 0 si aucune)", "type": "number", "required": True},
            {"key": "convention_collective", "label": "Convention collective applicable", "type": "text", "required": False},
            {"key": "lieu_travail", "label": "Lieu de travail principal", "type": "text", "required": True},
            {"key": "teletravail", "label": "Jours de télétravail/semaine", "type": "number", "required": False},
        ],
        "prompt_template": "cdi_v1",
    },
    "rupture_conventionnelle": {
        "name": "Protocole de Rupture Conventionnelle",
        "category": "rh",
        "description": "Document d'information et de protocole pour la rupture conventionnelle homologuée.",
        "version": "1.0",
        "required_plan": "pro",
        "fields": [
            {"key": "employeur_nom", "label": "Nom de l'employeur", "type": "text", "required": True},
            {"key": "employe_nom", "label": "Nom du salarié", "type": "text", "required": True},
            {"key": "date_premier_entretien", "label": "Date du 1er entretien prévu", "type": "date", "required": True},
            {"key": "date_rupture_souhaitee", "label": "Date de rupture souhaitée", "type": "date", "required": True},
            {"key": "salaire_brut_mensuel", "label": "Salaire brut mensuel (€)", "type": "number", "required": True},
            {"key": "anciennete_annees", "label": "Ancienneté (années)", "type": "number", "required": True},
            {"key": "indemnite_proposee", "label": "Indemnité proposée (€)", "type": "number", "required": True},
        ],
        "prompt_template": "rupture_conventionnelle_v1",
    },
    "avertissement": {
        "name": "Lettre d'Avertissement Salarié",
        "category": "rh",
        "description": "Lettre d'avertissement disciplinaire conforme au Code du travail.",
        "version": "1.0",
        "required_plan": "pro",
        "fields": [
            {"key": "employeur_nom", "label": "Nom de l'employeur", "type": "text", "required": True},
            {"key": "employe_nom", "label": "Nom du salarié", "type": "text", "required": True},
            {"key": "date_faits", "label": "Date des faits reprochés", "type": "date", "required": True},
            {"key": "description_faits", "label": "Description précise des faits", "type": "textarea", "required": True},
            {"key": "obligations_violees", "label": "Obligations contractuelles non respectées", "type": "textarea", "required": True},
        ],
        "prompt_template": "avertissement_v1",
    },
}

# ─── Prompts de génération par template ──────────────────────────────────────

GENERATION_PROMPTS = {
    "prestation_services_v1": """Tu es un expert en droit des contrats français. Génère un contrat de prestation de services professionnel et complet.

DONNÉES :
{form_data_formatted}

INSTRUCTIONS :
- Rédige un contrat complet de 800-1200 mots en français juridique accessible
- Inclus OBLIGATOIREMENT : objet, durée, prix, modalités de paiement, obligations des parties, responsabilité (PLAFONNÉE au montant du contrat), confidentialité si demandée, résiliation, droit applicable (droit français), juridiction compétente
- Pour la responsabilité : inclure "la responsabilité du prestataire est limitée au montant total du présent contrat" (Article 1231-3 Code civil)
- Inclure les pénalités de retard légales (Article L441-10 Code de Commerce)
- Format : structure claire avec articles numérotés
- Retourne UNIQUEMENT le texte du contrat, sans commentaires

{disclaimer}""",

    "cgv_v1": """Génère des Conditions Générales de Vente (CGV) conformes au droit français.

DONNÉES :
{form_data_formatted}

INCLURE OBLIGATOIREMENT :
- Article 1 : Objet et champ d'application
- Article 2 : Prix et modalités de paiement (avec pénalités de retard légales : 3× le taux d'intérêt légal, minimum 40€ d'indemnité)
- Article 3 : Livraison / exécution
- Article 4 : Droit de rétractation (14 jours si B2C)
- Article 5 : Garanties légales (conformité et vices cachés pour produits)
- Article 6 : Responsabilité
- Article 7 : Données personnelles (RGPD)
- Article 8 : Règlement des litiges et droit applicable
- Article 9 : Dispositions diverses

{disclaimer}""",

    "nda_v1": """Génère un accord de confidentialité (NDA) conforme au droit français.

DONNÉES :
{form_data_formatted}

INCLURE : définition des informations confidentielles, obligations des parties, exceptions (informations déjà publiques), durée, sanctions en cas de violation, droit applicable français, juridiction

{disclaimer}""",

    "relance_1_v1": """Génère une lettre de relance amiable de premier niveau, courtoise mais ferme.

DONNÉES :
{form_data_formatted}

STYLE : Professionnel et courtois. Rappeler les faits (facture, montant, échéance). Demander le règlement dans un délai de 8 jours. Mentionner qu'une suite pourra être donnée en l'absence de réponse. NE PAS menacer de procédure judiciaire (c'est la relance niveau 1).

{disclaimer}""",

    "mise_en_demeure_v1": """Génère une mise en demeure formelle conforme au droit français.

DONNÉES :
{form_data_formatted}

STYLE : Formel et sans ambiguïté. Mentionner : rappel des faits, base légale (Article 1231-1 du Code civil pour inexécution contractuelle), délai impératif de {delai} jours, et conséquences en cas de non-exécution (saisine du tribunal compétent). Indiquer que la lettre vaut mise en demeure au sens de l'article 1344 du Code civil.

{disclaimer}""",

    "cdi_v1": """Génère un contrat de travail CDI conforme au Code du travail français.

DONNÉES :
{form_data_formatted}

INCLURE OBLIGATOIREMENT :
- Identification complète des parties
- Article sur la période d'essai (si applicable, avec durées légales max)
- Classification et coefficient (selon convention collective si précisée)
- Rémunération brute mensuelle
- Durée du travail et organisation du temps de travail
- Lieu de travail et clause de mobilité (si applicable)
- Clause de télétravail (si applicable)
- Obligations du salarié (confidentialité, non-concurrence si applicable)
- Conditions de rupture du contrat
- Droit applicable : Code du travail français, Convention collective mentionnée

Références légales : L1221-1, L1237-19, L3121-27, L1237-1 et suivants Code du travail

{disclaimer}""",

    "rupture_conventionnelle_v1": """Génère un document d'information sur la procédure de rupture conventionnelle et un protocole d'accord.

DONNÉES :
{form_data_formatted}

INCLURE : rappel de la procédure (2 entretiens min, délai de rétractation 15 jours calendaires), calcul de l'indemnité légale minimum, montant proposé, calendrier, signature des deux parties, information sur l'homologation DREETS.

Base légale : Articles L1237-11 à L1237-16 du Code du travail

CALCUL INDEMNITÉ MINIMUM : {anciennete} × {salaire} × 1/4 (pour les 10 premières années)

{disclaimer}""",

    "avertissement_v1": """Génère une lettre d'avertissement disciplinaire conforme au Code du travail.

DONNÉES :
{form_data_formatted}

INCLURE : date et nature précise des faits, rappel des obligations contractuelles non respectées, avertissement formel, mention que cet avertissement restera dans le dossier disciplinaire pour 3 ans, demande d'explication optionnelle (recommandé pour se prémunir de toute contestation).

Base légale : Article L1331-1 du Code du travail (avertissement = sanction mineure ne nécessitant pas d'entretien préalable)

{disclaimer}""",
}


# ─── Service de génération ────────────────────────────────────────────────────

class DocumentGenerationService:

    def __init__(self, db: AsyncSession):
        self.db = db

    def template_exists(self, template_key: str) -> bool:
        return template_key in DOCUMENT_TEMPLATES

    def get_template_category(self, template_key: str) -> tuple[str, str]:
        tpl = DOCUMENT_TEMPLATES.get(template_key, {})
        return tpl.get("category", "other"), template_key

    async def generate_document(
        self,
        document_id: str,
        template_key: str,
        form_data: dict,
        org_name: str,
    ) -> None:
        """Génère un document et sauvegarde le résultat en BDD + S3."""
        logger.info("generation_started", doc_id=document_id, template=template_key)

        document = await self.db.get(Document, document_id)
        if not document:
            return

        document.status = DocumentStatus.PROCESSING
        await self.db.flush()

        try:
            template = DOCUMENT_TEMPLATES[template_key]
            prompt_key = template["prompt_template"]
            prompt_template = GENERATION_PROMPTS[prompt_key]

            # Construire le disclaimer daté
            disclaimer = LEGAL_DISCLAIMER.format(
                date=datetime.now(timezone.utc).strftime("%d/%m/%Y"),
                version=template["version"],
            )

            # Formater les données du formulaire
            form_data_formatted = "\n".join(
                f"- {k}: {v}" for k, v in form_data.items() if v is not None and v != ""
            )

            # Construire le prompt final
            prompt = prompt_template.format(
                form_data_formatted=form_data_formatted,
                disclaimer=disclaimer,
                **form_data,  # Permet d'injecter des champs individuels dans le prompt
            )

            # Appeler Mistral
            generated_text = await self._call_mistral_generation(prompt)

            # Générer le DOCX
            docx_bytes = await self._create_docx(
                content=generated_text,
                title=document.title,
                disclaimer=disclaimer,
            )

            # Uploader vers S3
            from app.services.documents.storage_service import StorageService
            storage_svc = StorageService()
            file_path = await storage_svc.upload_generated_document(
                content=docx_bytes,
                document_id=document_id,
                organization_id=document.organization_id,
                extension="docx",
            )

            # Mettre à jour le document
            document.status = DocumentStatus.COMPLETED
            document.file_path = file_path
            document.template_version = template["version"]
            document.generated_content = {
                "template_key": template_key,
                "form_data": form_data,
                "generated_text": generated_text[:5000],  # Stocker en BDD pour aperçu
            }
            await self.db.flush()

            logger.info("generation_completed", doc_id=document_id)

        except Exception as e:
            logger.error("generation_failed", doc_id=document_id, error=str(e))
            document.status = DocumentStatus.FAILED
            document.error_message = str(e)[:500]
            await self.db.flush()
            raise

    async def _call_mistral_generation(self, prompt: str) -> str:
        """Appelle Mistral pour la génération du document."""
        from mistralai import Mistral
        import asyncio

        client = Mistral(api_key=settings.MISTRAL_API_KEY)

        system_prompt = (
            "Tu es un expert juridique français spécialisé dans la rédaction de documents pour PME. "
            "Tu génères des documents professionnels, complets et conformes au droit français en vigueur. "
            "Utilise un langage juridique précis mais accessible. "
            "Ne génère que le texte du document demandé, sans introduction ni commentaire."
        )

        for attempt in range(3):
            try:
                response = client.chat.complete(
                    model=settings.MISTRAL_MODEL,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": prompt},
                    ],
                    max_tokens=3000,
                    temperature=0.05,  # Très bas pour les documents légaux
                )
                return response.choices[0].message.content
            except Exception as e:
                if attempt == 2:
                    raise
                await asyncio.sleep(2 ** attempt)

    async def _create_docx(self, content: str, title: str, disclaimer: str) -> bytes:
        """Crée un fichier DOCX formaté à partir du texte généré."""
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()

        # Titre du document
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date de génération
        date_para = doc.add_paragraph(
            f"Document généré le {datetime.now(timezone.utc).strftime('%d/%m/%Y')} par JurisIA"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(9)
        date_para.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph("")  # Espace

        # Contenu du document (préserver les sauts de ligne)
        for line in content.split("\n"):
            if line.strip():
                if line.strip().startswith("Article") or line.strip().startswith("ARTICLE"):
                    doc.add_heading(line.strip(), level=2)
                else:
                    doc.add_paragraph(line.strip())

        doc.add_paragraph("")  # Espace avant disclaimer

        # Disclaimer légal
        disclaimer_para = doc.add_paragraph(disclaimer)
        disclaimer_para.runs[0].font.size = Pt(8)
        disclaimer_para.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        disclaimer_para.runs[0].font.italic = True

        # Sauvegarder en mémoire
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
